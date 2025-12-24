import sys
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
# This script reads a text file with accent markers and generates a Word document
# with proper formatting for Vedic accent marks using the Adishila Vedic font.  
# Now deprecated in favor of generate_and_compile_latex.py


# Accent definitions using normal parentheses
accent_map = {
    "(1)": {"char": "\u0951", "size": 24, "bold": True},   # Swarita
    "(2)": {"char": "\u1CD2", "size": 28, "bold": True},   # Anudatta
    "(3)": {"char": "\u1CF8", "size": 24, "bold": True},   # Kampa
    "(4)": {"char": "\u1CF9", "size": 24, "bold": True},   # Trikampa
}

# Check for input file argument
if len(sys.argv) < 2:
    print("Usage: python format_accents.py <input_file.txt>")
    sys.exit(1)

input_file = sys.argv[1]

# Load input text
with open(input_file, "r", encoding="utf-8") as f:
    text = f.read()

# Create Word document
doc = Document()
para = doc.add_paragraph()

i = 0
while i < len(text):
    # Check for accent marker like (1), (2), etc.
    if text[i] == "(" and i + 2 < len(text) and text[i+2] == ")":
        code = text[i:i+3]
        if code in accent_map:
            accent = accent_map[code]
            run = para.add_run(accent["char"])
            run.font.name = "Adishila Vedic"
            run.font.size = Pt(accent["size"])
            run.bold = accent["bold"]
            r = run._element
            r.rPr.rFonts.set(qn("w:eastAsia"), "Adishila Vedic")
            i += 3
            continue
    # Regular character
    run = para.add_run(text[i])
    run.font.name = "Adishila Vedic"
    run.font.size = Pt(16)
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), "Adishila Vedic")
    i += 1

# Save output
output_file = input_file.rsplit(".", 1)[0] + "_formatted.docx"
doc.save(output_file)
print(f"Formatted document saved as: {output_file}")